#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Professional Word (.docx) Outline Generator Engine with Deep Navy / Slate styling.
Takes JSON outline data and produces a structured, beautifully formatted Word document.
"""

import os
import sys
import json
import argparse
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Palette Constants
COLOR_PRIMARY   = RGBColor(24, 43, 73)      # Deep Navy #182B49
COLOR_SECONDARY = RGBColor(41, 128, 185)  # Ocean Blue #2980B9
COLOR_ACCENT    = RGBColor(184, 134, 11)     # Gold #B8860B
COLOR_TEXT      = RGBColor(44, 62, 80)         # Charcoal #2C3E50
COLOR_BG_LIGHT  = "F4F6F9"
COLOR_BG_HEADER = "182B49"

def set_cell_bg(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_outline_document(data, output_docx_path):
    doc = docx.Document()
    
    # Page Margins
    for s in doc.sections:
        s.page_height = Inches(11.69)
        s.page_width = Inches(8.27)
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run(data.get("title", "影片大綱與時間戳紀錄"))
    r.font.name = "微軟正黑體"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = COLOR_PRIMARY
    
    # Subtitle
    if data.get("subtitle"):
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(12)
        r = p_sub.add_run(data["subtitle"])
        r.font.name = "微軟正黑體"
        r.font.size = Pt(12)
        r.font.color.rgb = COLOR_SECONDARY
        
    # Metadata Table
    meta = data.get("metadata", {})
    if meta:
        meta_table = doc.add_table(rows=2, cols=4)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_items = list(meta.items())
        
        for r_idx in range(2):
            for c_idx in range(2):
                idx = r_idx * 2 + c_idx
                if idx < len(meta_items):
                    k, v = meta_items[idx]
                    cell_k = meta_table.cell(r_idx, c_idx * 2)
                    cell_v = meta_table.cell(r_idx, c_idx * 2 + 1)
                    
                    cell_k.text = k
                    set_cell_bg(cell_k, "EBF3FA")
                    cell_k.paragraphs[0].runs[0].font.name = "微軟正黑體"
                    cell_k.paragraphs[0].runs[0].font.size = Pt(9)
                    cell_k.paragraphs[0].runs[0].font.bold = True
                    cell_k.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
                    set_cell_margins(cell_k, 70, 70, 90, 90)
                    
                    cell_v.text = str(v)
                    set_cell_bg(cell_v, "FAFAFA")
                    cell_v.paragraphs[0].runs[0].font.name = "微軟正黑體"
                    cell_v.paragraphs[0].runs[0].font.size = Pt(9)
                    cell_v.paragraphs[0].runs[0].font.color.rgb = COLOR_TEXT
                    set_cell_margins(cell_v, 70, 70, 90, 90)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        
    # Summary Table with Timestamps
    chapters = data.get("chapters", [])
    if chapters:
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(14)
        h1.paragraph_format.space_after = Pt(6)
        r = h1.add_run("【章節導覽與精確時間戳總表】")
        r.font.name = "微軟正黑體"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        
        sum_table = doc.add_table(rows=1, cols=4)
        sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_widths = [Inches(1.2), Inches(1.1), Inches(2.1), Inches(2.27)]
        
        hdr = ["時間戳", "章節", "核心主題", "重點摘要"]
        for i, h in enumerate(hdr):
            cell = sum_table.cell(0, i)
            cell.text = h
            cell.width = c_widths[i]
            set_cell_bg(cell, COLOR_BG_HEADER)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            p.runs[0].font.name = "微軟正黑體"
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
        for r_idx, ch in enumerate(chapters):
            row = sum_table.add_row().cells
            bg = "FFFFFF" if r_idx % 2 == 0 else COLOR_BG_LIGHT
            values = [ch.get("timestamp", ""), ch.get("unit", ""), ch.get("title", ""), ch.get("summary", "")]
            for c_idx, val in enumerate(values):
                row[c_idx].text = val
                row[c_idx].width = c_widths[c_idx]
                set_cell_bg(row[c_idx], bg)
                set_cell_margins(row[c_idx], 60, 60, 90, 90)
                p = row[c_idx].paragraphs[0]
                p.runs[0].font.name = "微軟正黑體"
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = COLOR_SECONDARY if c_idx == 0 else COLOR_TEXT
                if c_idx == 0:
                    p.runs[0].font.bold = True
                    
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
        
    # Detailed Sections
    sections = data.get("sections", [])
    if sections:
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(14)
        h1.paragraph_format.space_after = Pt(6)
        r = h1.add_run("【詳細內容剖析與重點整理】")
        r.font.name = "微軟正黑體"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        
        for sec in sections:
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(10)
            p_sec.paragraph_format.space_after = Pt(3)
            r = p_sec.add_run(sec.get("heading", ""))
            r.font.name = "微軟正黑體"
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = COLOR_SECONDARY
            
            for item in sec.get("bullets", []):
                p_b = doc.add_paragraph(style='List Bullet')
                p_b.paragraph_format.space_after = Pt(2)
                if isinstance(item, dict):
                    prefix = item.get("prefix", "")
                    content = item.get("text", "")
                    if prefix:
                        rp = p_b.add_run(prefix + " ")
                        rp.font.name = "微軟正黑體"
                        rp.font.size = Pt(9.5)
                        rp.font.bold = True
                        rp.font.color.rgb = COLOR_PRIMARY
                    rc = p_b.add_run(content)
                    rc.font.name = "微軟正黑體"
                    rc.font.size = Pt(9.5)
                    rc.font.color.rgb = COLOR_TEXT
                else:
                    rc = p_b.add_run(str(item))
                    rc.font.name = "微軟正黑體"
                    rc.font.size = Pt(9.5)
                    rc.font.color.rgb = COLOR_TEXT

    doc.save(output_docx_path)
    print(f"Word Document successfully saved to: {output_docx_path}")
    return True

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate styled Word document from JSON outline data.")
    parser.add_argument('--input', required=True, help="Input JSON outline data file")
    parser.add_argument('--output', required=True, help="Output .docx file path")
    args = parser.parse_args()
    
    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)
    create_outline_document(data, args.output)
